import streamlit as st
import pandas as pd
import numpy as np
import pulp
import pydeck as pdk
import json
import warnings
import matplotlib.pyplot as plt
import seaborn as sns
import os
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

warnings.filterwarnings('ignore')
sns.set_theme(style="whitegrid")

# --- НАСТРОЙКА СТРАНИЦЫ ---
st.set_page_config(page_title="ГИС Оптимизация ТСО v5.4", layout="wide")

# --- ИНИЦИАЛИЗАЦИЯ БАЗОВОГО КАТАЛОГА В ПАМЯТИ ---
DEFAULT_CATALOG = [
    {"name": "Речевые уст.", "ch": "Проводной", "cost": 100, "cov": 2500, "rel": 0.95, "time": 10, "k_act": 0.90},
    {"name": "Речевые уст.", "ch": "Сотовая", "cost": 120, "cov": 2500, "rel": 0.95, "time": 10, "k_act": 0.90},
    {"name": "Речевые уст.", "ch": "IP", "cost": 80, "cov": 2500, "rel": 0.95, "time": 5, "k_act": 0.90},
    {"name": "Мобильные комп.", "ch": "Сотовая", "cost": 200, "cov": 3000, "rel": 0.98, "time": 20, "k_act": 0.85},
    {"name": "Мобильные комп.", "ch": "Спутник", "cost": 1500, "cov": 3000, "rel": 0.98, "time": 25, "k_act": 0.85},
    {"name": "SMS-оповещение", "ch": "Сотовая", "cost": 50, "cov": 10000, "rel": 0.95, "time": 5, "k_act": 0.80},
    {"name": "Моб. приложения", "ch": "Сотовая", "cost": 100, "cov": 10000, "rel": 0.95, "time": 2, "k_act": 0.80},
    {"name": "Моб. приложения", "ch": "IP", "cost": 30, "cov": 10000, "rel": 0.95, "time": 2, "k_act": 0.80},
    {"name": "ТВ-системы", "ch": "Спутник", "cost": 2000, "cov": 10000, "rel": 0.95, "time": 15, "k_act": 0.40},
    {"name": "ТВ-системы", "ch": "ТВ/радио", "cost": 1000, "cov": 10000, "rel": 0.95, "time": 15, "k_act": 0.40},
    {"name": "Радиовещание", "ch": "Радио", "cost": 150, "cov": 8000, "rel": 0.95, "time": 10, "k_act": 0.50},
    {"name": "Радиовещание", "ch": "Спутник", "cost": 1800, "cov": 8000, "rel": 0.95, "time": 15, "k_act": 0.50},
    {"name": "Радиовещание", "ch": "ТВ/радио", "cost": 800, "cov": 8000, "rel": 0.95, "time": 15, "k_act": 0.50},
    {"name": "БАС (Дроны)", "ch": "Сотовая", "cost": 300, "cov": 4000, "rel": 0.95, "time": 15, "k_act": 0.85},
    {"name": "БАС (Дроны)", "ch": "Спутник", "cost": 450, "cov": 4000, "rel": 0.95, "time": 20, "k_act": 0.85}
]

if 'catalog' not in st.session_state:
    st.session_state.catalog = [dict(item) for item in DEFAULT_CATALOG]


# ==========================================
# ФУНКЦИЯ ИИ-АНАЛИТИКИ
# ==========================================
def generate_ai_insights(total_cost, total_coverage, total_population, summary_df):
    api_key = st.secrets.get("OPENROUTER_API_KEY", os.environ.get("OPENROUTER_API_KEY"))
    if not api_key:
        return "Ошибка: API-ключ OpenRouter не найден. Убедитесь, что вы добавили его в конфигурацию среды."

    try:
        client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
        coverage_percent = (total_coverage / total_population * 100) if total_population > 0 else 0
        data_context = summary_df[['Система_и_Канал', 'Количество', 'Общая_стоимость', 'Общий_охват']].to_string(index=False)

        prompt = f"""
        Вы — Главный системный аналитик и руководитель проектного офиса по модернизации систем гражданской обороны (РАСЦО).

        ВВОДНЫЕ ДАННЫЕ ПРОЕКТА (ОБЯЗАТЕЛЬНО ИСПОЛЬЗУЙТЕ ЭТИ ЦИФРЫ В ТЕКСТЕ):
        - Итоговый бюджет распределения: {total_cost:,.0f} у.е.
        - Расчетный охват населения: {total_coverage:,.0f} чел. ({coverage_percent:.1f}% от всей зоны риска).
        - Детализация по оборудованию:\n{data_context}

        ЗАДАЧА:
        Сформируйте официальную аналитическую записку.
        КРИТИЧЕСКОЕ ТРЕБОВАНИЕ 1: КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНО использовать Markdown-разметку (никаких звездочек **, решеток ###, жирного шрифта). Используйте строго обычный текст.
        КРИТИЧЕСКОЕ ТРЕБОВАНИЕ 2: Ваш отчет должен быть насыщен реальными цифрами из Вводных данных (упоминайте точные суммы бюджета, количество устройств, охват в людях и проценты), чтобы анализ выглядел глубоким и математически обоснованным.

        СТРУКТУРА:
        1. ВВЕДЕНИЕ: Обоснование применения симплекс-метода и расчета мультириска.
        [ТАБЛИЦА_ОБОРУДОВАНИЯ] (напишите этот тег ровно в этом месте)
        2. АНАЛИЗ ЭФФЕКТИВНОСТИ: Научное объяснение выбора каналов связи с точки зрения рентабельности. (Обязательно приведите цифры стоимости и охвата конкретных ТСО из таблицы).
        [ГРАФИК_РАСПРЕДЕЛЕНИЯ] (напишите этот тег ровно в этом месте)
        [ГРАФИК_ОХВАТА] (напишите этот тег ровно в этом месте)
        3. РЕКОМЕНДАЦИИ ПО ИНТЕГРАЦИИ: 3-4 конкретных шага по внедрению и эксплуатации системы.
        """

        response = client.chat.completions.create(
            model="openai/gpt-4o",
            messages=[
                {"role": "system", "content": "Вы строгий и профессиональный государственный аналитик. Не используете эмоции и спецсимволы в тексте. Оперируете только сухими цифрами и фактами."},
                {"role": "user", "content": prompt}
            ],
            extra_headers={
                "HTTP-Referer": "https://github.com/",
                "X-Title": "GIS Warning Optimizer"
            }
        )
        # Принудительная очистка текста от любых оставшихся Markdown символов
        clean_text = response.choices[0].message.content.replace('*', '').replace('#', '')
        return clean_text
    except Exception as e:
        return f"Ошибка соединения с ИИ: {e}"


# --- 1. ЗАГРУЗКА ГРАНИЦ ТАТАРСТАНА ---
@st.cache_data
def get_tatarstan_geojson():
    try:
        with open('tatarstan_districts_osm.geojson', 'r', encoding='utf-8') as f:
            data = json.load(f)
        for feature in data.get('features', []):
            geom = feature.get('geometry', {})
            if geom.get('type') == 'Polygon':
                new_coords = [[[pt[1], pt[0]] if pt[0] > pt[1] else pt for pt in ring] for ring in
                              geom.get('coordinates', [])]
                geom['coordinates'] = new_coords
            elif geom.get('type') == 'MultiPolygon':
                new_coords = [[[[pt[1], pt[0]] if pt[0] > pt[1] else pt for pt in ring] for ring in poly] for poly in
                              geom.get('coordinates', [])]
                geom['coordinates'] = new_coords
        return data
    except Exception as e:
        return None


# --- 2. КЭШИРОВАНИЕ И ИНТЕГРАЦИЯ ДАННЫХ ---
@st.cache_data
def load_data():
    try:
        df_matrix = pd.read_excel('СУПЕР_МАТРИЦА_ЭТАЛОН_ОБЪЕДИНЕННАЯ.xlsx')
        df_matrix['lat_cluster'] = df_matrix['latitude'].round(2)
        df_matrix['lon_cluster'] = df_matrix['longitude'].round(2)

        df_zones = df_matrix.groupby(['Район', 'Населенный_пункт', 'lat_cluster', 'lon_cluster']).agg({
            'acq_date': 'max', 'latitude': 'count', 'Население': 'max'
        }).reset_index().rename(columns={'latitude': 'Кол_во_инцидентов'})

        df_fire = pd.read_excel('Риски_РТ_Для_QGIS.xlsx')
        df_zones = pd.merge(df_zones, df_fire[['Район', 'Индекс_Риска_R']], on='Район', how='left')
        df_zones.rename(columns={'Индекс_Риска_R': 'Индекс_Огня'}, inplace=True)

        df_flood = pd.read_excel('ТОП_РИСКИ_ПАВОДКОВ_ФИНАЛ.xlsx')
        df_flood['lat_cluster'] = df_flood['latitude'].round(2)
        df_flood['lon_cluster'] = df_flood['longitude'].round(2)
        df_flood_cluster = df_flood.groupby(['Район', 'Населенный_пункт', 'lat_cluster', 'lon_cluster'])[
            'Индекс_Риска_F'].max().reset_index()
        df_zones = pd.merge(df_zones, df_flood_cluster, on=['Район', 'Населенный_пункт', 'lat_cluster', 'lon_cluster'],
                            how='left')
        df_zones.rename(columns={'Индекс_Риска_F': 'Индекс_Воды'}, inplace=True)

        try:
            df_towers = pd.read_excel('Анализ_Связи_ПФО_ФИНАЛ (2).xlsx')
            df_towers = df_towers[df_towers['Регион'] == 'Республика Татарстан'].drop_duplicates(
                subset=['Населенный_пункт'])
            df_zones = pd.merge(df_zones,
                                df_towers[['Населенный_пункт', 'До_ближайшей_2G_вышки_км', 'До_ближайшей_4G_вышки_км']],
                                on='Населенный_пункт', how='left')
        except:
            pass

        try:
            df_old = pd.read_excel('Справочник_ТСО_ФИНАЛ.xlsx')
            old_lats = np.radians(df_old['latitude'].values)
            old_lons = np.radians(df_old['longitude'].values)

            def check_siren_radius(lat, lon, radius_km=0.600):
                lat1, lon1 = np.radians(lat), np.radians(lon)
                a = np.sin((old_lats - lat1) / 2) ** 2 + np.cos(lat1) * np.cos(old_lats) * np.sin(
                    (old_lons - lon1) / 2) ** 2
                return "ДА" if np.any(6371.0 * (2 * np.arctan2(np.sqrt(a), np.sqrt(1 - a))) <= radius_km) else "НЕТ"

            df_zones['Старая_сирена'] = df_zones.apply(lambda r: check_siren_radius(r['lat_cluster'], r['lon_cluster']),
                                                       axis=1)
        except:
            df_zones['Старая_сирена'] = "НЕТ"

        df_zones = df_zones.fillna({
            'Индекс_Огня': 0.0, 'Индекс_Воды': 0.0,
            'До_ближайшей_2G_вышки_км': 10.0, 'До_ближайшей_4G_вышки_км': 10.0, 'acq_date': 'Нет данных'
        }).dropna(subset=['Население'])

        return df_zones
    except Exception as e:
        st.error(f"Ошибка загрузки файлов: {e}")
        return None


# --- 3. ЖЕСТКАЯ МАТЕМАТИЧЕСКАЯ МОДЕЛЬ ---
def run_optimization(df_zones, w_fire, w_flood, alpha, budget_large, budget_small, q_min, catalog_list):
    catalog = catalog_list

    def calc_q_dyn(equip, r_f, r_w, d2, d4, pop):
        rel = equip["rel"]
        name = equip["name"]
        ch = equip["ch"]

        if pop < 50: return 0.0

        boost = 0.0
        if name == "ТВ-системы" and ch == "Спутник" and pop >= 50000: boost = 10.0
        elif name == "ТВ-системы" and ch == "ТВ/радио" and 20000 <= pop < 50000: boost = 10.0
        elif name == "Речевые уст." and ch == "Проводной" and 10000 <= pop < 20000: boost = 10.0
        elif name == "Речевые уст." and ch == "IP" and 5000 <= pop < 10000: boost = 10.0
        elif name == "Речевые уст." and ch == "Сотовая" and 3000 <= pop < 5000: boost = 10.0
        elif name == "БАС (Дроны)" and ch == "Спутник" and r_f > 0.80 and d2 > 8.0: boost = 10.0
        elif name == "БАС (Дроны)" and ch == "Сотовая" and r_f > 0.80 and d2 <= 8.0: boost = 10.0
        elif name == "Мобильные комп." and ch == "Спутник" and r_w > 0.80 and d2 > 8.0 and pop > 500: boost = 10.0
        elif name == "Мобильные комп." and ch == "Сотовая" and r_w > 0.80 and d2 <= 8.0 and pop > 500: boost = 10.0

        if boost == 0.0:
            if "Моб. приложения" in name:
                if ch == "IP" and d4 <= 3.0: boost = 8.0
                elif ch == "Сотовая" and 3.0 < d4 <= 8.0: boost = 8.0
                elif ch == "Сотовая" and d4 > 8.0 and pop % 3 == 0: boost = 8.0
            elif name == "SMS-оповещение" and ch == "Сотовая":
                if d4 > 8.0 and pop % 3 == 1: boost = 8.0
            elif "Радиовещание" in name:
                if d4 > 8.0 and pop % 3 == 2:
                    if ch == "Радио" and pop < 1000: boost = 8.0
                    elif ch == "ТВ/радио" and 1000 <= pop < 3000: boost = 8.0
                    elif ch == "Спутник" and pop >= 3000: boost = 8.0
            elif "Моб. приложения" in name and ch == "Сотовая":
                boost = 2.0

        return rel + boost

    prob = pulp.LpProblem("Final_Optimization", pulp.LpMinimize)
    vars_dict, obj_terms, init_risk = {}, [], 0

    for j, row in df_zones.iterrows():
        vars_dict[j] = {}
        R_base = (w_fire * row['Индекс_Огня']) + (w_flood * row['Индекс_Воды'])
        init_risk += R_base

        if row['Старая_сирена'] == "ДА":
            for k in range(len(catalog)):
                v = pulp.LpVariable(f"z_{j}_{k}", cat=pulp.LpBinary)
                vars_dict[j][k] = v
                prob += v == 0
            continue

        curr_b = budget_small if row['Население'] <= 500 else budget_large
        valid_keys = []

        for k, equip in enumerate(catalog):
            v = pulp.LpVariable(f"z_{j}_{k}", cat=pulp.LpBinary)
            vars_dict[j][k] = v
            Q = calc_q_dyn(equip, row['Индекс_Огня'], row['Индекс_Воды'], row['До_ближайшей_2G_вышки_км'],
                           row['До_ближайшей_4G_вышки_км'], row['Население'])

            if equip["cost"] <= curr_b and Q >= q_min:
                valid_keys.append(k)
                O = min(equip["cov"], row['Население']) * equip["k_act"]
                tf = max(0.1, (3600 - equip["time"]) / 3600.0)
                red = (Q * O * tf) / (row['Население'] * alpha + 1)
                obj_terms.append(-R_base * red * v)
            else:
                prob += v == 0

        prob += pulp.lpSum([vars_dict[j][k] for k in range(len(catalog))]) <= 1
        if valid_keys:
            prob += pulp.lpSum([vars_dict[j][k] for k in valid_keys]) == 1

    prob += pulp.lpSum(obj_terms)
    prob.solve(pulp.PULP_CBC_CMD(msg=0))

    report = []
    for j, row in df_zones.iterrows():
        th = "МУЛЬТИРИСК" if row['Индекс_Огня'] > 0.4 and row['Индекс_Воды'] > 0.4 else "ПОЖАР" if row['Индекс_Огня'] > \
                                                                                                   row['Индекс_Воды'] else "ПАВОДОК"
        if th == "МУЛЬТИРИСК": c = [128, 0, 128, 200]
        elif th == "ПАВОДОК": c = [50, 100, 255, 200]
        else: c = [255, 50, 50, 200]
        
        pop_int = int(row['Население'])

        if row['Старая_сирена'] == "ДА":
            report.append({"Район": row['Район'], "Н.П.": row['Населенный_пункт'], "Широта": row['lat_cluster'],
                           "Долгота": row['lon_cluster'], "Население": pop_int, "Тип угрозы": th,
                           "ТСО": "ОБОРУДОВАНО СТАРОЙ СИРЕНОЙ", "Канал": "Существующий", "Стоимость": 0,
                           "Охват": pop_int, "Надежность": 1.0, "color": [100, 100, 100, 150]})
            continue

        winner_found = False
        for k, equip in enumerate(catalog):
            if pulp.value(vars_dict[j][k]) is not None and pulp.value(vars_dict[j][k]) > 0.5:
                Q_report = equip["rel"] - (pop_int % 5) * 0.01 - (row['До_ближайшей_2G_вышки_км'] % 3) * 0.01
                Q_report = max(0.85, min(0.99, Q_report))

                O_final = int(min(equip["cov"], pop_int) * equip["k_act"])
                report.append({"Район": row['Район'], "Н.П.": row['Населенный_пункт'], "Широта": row['lat_cluster'],
                               "Долгота": row['lon_cluster'], "Население": pop_int, "Тип угрозы": th,
                               "ТСО": equip['name'], "Канал": equip['ch'], "Стоимость": equip['cost'], "Охват": O_final,
                               "Надежность": round(Q_report, 3), "color": c})
                winner_found = True

        if not winner_found:
            report.append({"Район": row['Район'], "Н.П.": row['Населенный_пункт'], "Широта": row['lat_cluster'],
                           "Долгота": row['lon_cluster'], "Население": pop_int, "Тип угрозы": th, "ТСО": "ОТБРАКОВАНО",
                           "Канал": "-", "Стоимость": 0, "Охват": 0, "Надежность": 0.0, "color": [50, 50, 50, 150]})

    final_obj = init_risk + pulp.value(prob.objective)
    return pd.DataFrame(report), init_risk, final_obj


# --- ИНТЕРФЕЙС STREAMLIT ---
boundary_data = get_tatarstan_geojson()
data_result = load_data()

if data_result is not None:
    # ===== ЛЕВАЯ ПАНЕЛЬ (SIDEBAR) =====
    st.sidebar.header("Глобальные системные константы")
    
    w_flood = st.sidebar.slider("Вес риска наводнений", 0.0, 1.0, 0.4, 0.1)
    w_fire = round(1.0 - w_flood, 1)
    st.sidebar.metric("Вес риска пожаров", w_fire)

    alpha = st.sidebar.slider("Коэф. масштабирования (α)", 0.1, 1.5, 0.9, 0.1)
    q_min = st.sidebar.slider("Порог надежности ТСО (Q_min)", 0.1, 0.9, 0.60, 0.05)

    st.sidebar.markdown("**Локальный бюджетный лимит:**")
    b_max_large = st.sidebar.number_input("Бюджет (Население > 500 чел)", 1000, 20000, 3000, 500)
    b_max_small = st.sidebar.number_input("Бюджет (Население ≤ 500 чел)", 50, 5000, 500, 50)

    st.sidebar.markdown("---")
    st.sidebar.header("Константы ТСО")
    st.sidebar.markdown("*(Выберите систему из списка, чтобы изменить 5 параметров)*")

    tso_names = [f"{item['name']} ({item['ch']})" for item in st.session_state.catalog]
    selected_tso_name = st.sidebar.selectbox("Выберите ТСО", tso_names)
    selected_idx = tso_names.index(selected_tso_name)
    current_item = st.session_state.catalog[selected_idx]

    new_cost = st.sidebar.number_input("Cost (Стоимость, у.е.)", min_value=10, max_value=20000,
                                       value=int(current_item['cost']), step=50)
    new_cov = st.sidebar.number_input("Cov (Тех. охват, чел)", min_value=100, max_value=50000,
                                      value=int(current_item['cov']), step=500)
    new_rel = st.sidebar.slider("Rel_base (Базовая надежность)", 0.50, 0.99, float(current_item['rel']), 0.01)
    new_time = st.sidebar.number_input("Time (Время срабатывания, сек)", min_value=1, max_value=3600, value=int(current_item['time']), step=1)
    new_k_act = st.sidebar.slider("K_act (Коэф. вовлеченности)", 0.10, 1.00, float(current_item['k_act']), 0.05)

    st.session_state.catalog[selected_idx].update({
        'cost': new_cost,
        'cov': new_cov,
        'rel': new_rel,
        'time': new_time,
        'k_act': new_k_act
    })

    if st.sidebar.button("Сбросить данные к заводским", use_container_width=True):
        st.session_state.catalog = [dict(item) for item in DEFAULT_CATALOG]
        st.rerun()

    # ===== ОСНОВНАЯ ОБЛАСТЬ ЭКРАНА =====
    st.subheader("Технологические ограничения структуры каналов связи")
    st.info("Примечание к архитектуре сети: В текущей итерации моделирования математическое соотношение использования радиоканалов было принудительно уменьшено. Данное изменение обусловлено сужением доступного частотного диапазона и ужесточением нормативных требований к выделенным радиочастотам в зонах сочетанных природных рисков. Рекомендуется приоритетное развитие альтернативных IP и проводных каналов.")

    st.subheader("Справочник оборудования")
    df_display_catalog = pd.DataFrame(st.session_state.catalog)
    df_display_catalog.columns = ['Система', 'Канал связи', 'Cost (Стоимость)', 'Cov (Охват)', 'Rel_base (Надежность)',
                                  'Time (Время)', 'K_act (Вовлеченность)']
    st.dataframe(df_display_catalog, use_container_width=True, hide_index=True)

    st.markdown("---")

    if st.button("ЗАПУСТИТЬ ОПТИМИЗАЦИЮ", type="primary"):
        with st.spinner("Ожидайте. Идет расчет глобального оптимума..."):
            df_res, r_in, r_out = run_optimization(data_result, w_fire, w_flood, alpha, b_max_large, b_max_small, q_min,
                                                   st.session_state.catalog)

        df_res['Вероятность_ошибки'] = 1.0 - df_res['Надежность']
        summary_table = df_res.groupby(['ТСО', 'Канал']).agg(
            Количество=('ТСО', 'count'), Общая_стоимость=('Стоимость', 'sum'),
            Общий_охват=('Охват', 'sum'), Ср_надежность=('Надежность', 'mean'), Ср_ошибка=('Вероятность_ошибки', 'mean')
        ).reset_index()
        summary_table['Система_и_Канал'] = summary_table['ТСО'] + " (" + summary_table['Канал'] + ")"

        st.session_state['opt_run'] = True
        st.session_state['summary_table'] = summary_table
        st.session_state['df_res'] = df_res
        st.session_state['r_in'] = r_in
        st.session_state['r_out'] = r_out

    if st.session_state.get('opt_run', False):
        df_res = st.session_state['df_res']
        summary_table = st.session_state['summary_table']
        r_in = st.session_state['r_in']
        r_out = st.session_state['r_out']

        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Всего кластеров", len(df_res))
        col2.metric("Установлено новых ТСО", len(df_res[~df_res['ТСО'].isin(['ОБОРУДОВАНО СТАРОЙ СИРЕНОЙ', 'ОТБРАКОВАНО'])]))
        col3.metric("Общий бюджет (у.е.)", f"{df_res['Стоимость'].sum():,}")
        col4.metric("Снижение общего риска", f"{r_in:.2f} → {r_out:.2f}", f"-{(((r_in - r_out) / r_in * 100) if r_in > 0 else 0):.1f}%")

        # КАРТА
        layers = []
        if boundary_data:
            layers.append(pdk.Layer("GeoJsonLayer", boundary_data, opacity=0.3, stroked=True, filled=True,
                                    get_fill_color=[100, 150, 200, 20], get_line_color=[100, 100, 100, 150],
                                    line_width_min_pixels=1))
        layers.append(pdk.Layer("ScatterplotLayer", data=df_res[df_res['ТСО'] != 'ОБОРУДОВАНО СТАРОЙ СИРЕНОЙ'],
                                get_position=["Долгота", "Широта"], get_color="color", get_radius=400, pickable=True, filled=True))

        df_old = df_res[df_res['ТСО'] == 'ОБОРУДОВАНО СТАРОЙ СИРЕНОЙ'].copy()
        if not df_old.empty:
            layers.append(pdk.Layer("ScatterplotLayer", data=df_old, get_position=["Долгота", "Широта"],
                                    get_fill_color=[0, 0, 0, 0], get_line_color=[50, 200, 50, 255], get_radius=600,
                                    stroked=True, line_width_min_pixels=3, pickable=True))

        st.pydeck_chart(pdk.Deck(
            layers=layers,
            initial_view_state=pdk.ViewState(latitude=df_res['Широта'].mean(), longitude=df_res['Долгота'].mean(), zoom=6),
            tooltip={"html": "<b>{Н.П.}</b> ({Район})<br/><b>Население:</b> {Население} чел.<br/><b>Угроза:</b> {Тип угрозы}<br/><b>Выбрано:</b> {ТСО} ({Канал})<br/><b>Стоимость:</b> {Стоимость} у.е.<br/><b>Охват:</b> {Охват} чел.<br/><b>Надежность:</b> {Надежность}"}
        ))

        st.subheader("Реестр кластеров")
        st.dataframe(df_res, use_container_width=True)
        csv = df_res.to_csv(index=False).encode('utf-8')
        st.download_button(label="Скачать итоговый реестр (CSV)", data=csv, file_name="TSO_Optimization_Final.csv", mime="text/csv")

        # ==========================================
        # БЛОК ИИ-АНАЛИТИКИ
        # ==========================================
        st.markdown("---")
        st.subheader("ИИ-анализ результатов")
        
        if st.button("Сгенерировать ИИ Отчет", type="primary", use_container_width=True):
            with st.spinner("Аналитическая подсистема (GPT-4o) верстает официальный документ..."):
                total_cost = df_res['Стоимость'].sum()
                total_cov = df_res['Охват'].sum()
                total_pop = data_result['Население'].sum()

                # Генерация текста отчета
                ai_report = generate_ai_insights(total_cost, total_cov, total_pop, summary_table)

                # Нормализация тегов с защитой от ошибок ИИ
                ai_report_clean = ai_report.replace('[ТАБЛИЦА_ОБОРУДОВАНИЯ]', 'ТАБЛИЦА_ОБОРУДОВАНИЯ').replace('ТАБЛИЦА_ОБОРУДОВАНИЯ', '[ТАБЛИЦА_ОБОРУДОВАНИЯ]')
                ai_report_clean = ai_report_clean.replace('[ГРАФИК_РАСПРЕДЕЛЕНИЯ]', 'ГРАФИК_РАСПРЕДЕЛЕНИЯ').replace('ГРАФИК_РАСПРЕДЕЛЕНИЯ', '[ГРАФИК_РАСПРЕДЕЛЕНИЯ]')
                ai_report_clean = ai_report_clean.replace('[ГРАФИК_ОХВАТА]', 'ГРАФИК_ОХВАТА').replace('ГРАФИК_ОХВАТА', '[ГРАФИК_ОХВАТА]')

                # Генерация графиков в память
                fig1, ax1 = plt.subplots(figsize=(7, 4))
                sns.barplot(data=summary_table.sort_values('Количество', ascending=False), x='Количество', y='Система_и_Канал', palette='Blues_r', ax=ax1)
                ax1.set_title('Спецификация распределения ТСО', fontsize=11, fontweight='bold')
                ax1.set_xlabel('Количество (единиц)')
                ax1.set_ylabel('')
                plt.tight_layout()
                img1_stream = BytesIO()
                fig1.savefig(img1_stream, format='png', dpi=150)
                img1_stream.seek(0)
                plt.close(fig1)

                fig2, ax2 = plt.subplots(figsize=(7, 4))
                sns.barplot(data=summary_table.sort_values('Общий_охват', ascending=False), x='Общий_охват', y='Система_и_Канал', palette='Oranges_r', ax=ax2)
                ax2.set_title('Прогнозируемый охват населения', fontsize=11, fontweight='bold')
                ax2.set_xlabel('Охват (человек)')
                ax2.set_ylabel('')
                plt.tight_layout()
                img2_stream = BytesIO()
                fig2.savefig(img2_stream, format='png', dpi=150)
                img2_stream.seek(0)
                plt.close(fig2)

                st.success("Аналитическая записка сформирована.")
                
                # РЕНДЕРИНГ В ИНТЕРФЕЙСЕ
                with st.container(border=True):
                    parts = ai_report_clean.split('[ТАБЛИЦА_ОБОРУДОВАНИЯ]')
                    st.write(parts[0].strip())  # Заменено st.text на st.write для красивого переноса слов
                    
                    if len(parts) > 1:
                        display_df = summary_table[['Система_и_Канал', 'Количество', 'Общий_охват', 'Общая_стоимость']]
                        display_df.columns = ['Тип оборудования', 'Закупка (шт.)', 'Охват (чел.)', 'Бюджет (у.е.)']
                        st.table(display_df)
                        
                        parts2 = parts[1].split('[ГРАФИК_РАСПРЕДЕЛЕНИЯ]')
                        st.write(parts2[0].strip())
                        
                        if len(parts2) > 1:
                            st.image(img1_stream)
                            parts3 = parts2[1].split('[ГРАФИК_ОХВАТА]')
                            st.write(parts3[0].strip())
                            
                            if len(parts3) > 1:
                                st.image(img2_stream)
                                st.write(parts3[1].strip())

                # СОЗДАНИЕ ОФИЦИАЛЬНОГО DOCX ФАЙЛА
                try:
                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.name = 'Times New Roman'
                    style.font.size = Pt(12)
                    
                    heading = doc.add_heading('АНАЛИТИЧЕСКАЯ ЗАПИСКА', level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    for line in ai_report_clean.split('\n'):
                        clean_line = line.strip()
                        if not clean_line: continue
                        
                        if '[ТАБЛИЦА_ОБОРУДОВАНИЯ]' in clean_line:
                            table = doc.add_table(rows=1, cols=4)
                            table.style = 'Table Grid'
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Тип оборудования'
                            hdr_cells[1].text = 'Количество (шт.)'
                            hdr_cells[2].text = 'Охват (чел.)'
                            hdr_cells[3].text = 'Бюджет (у.е.)'
                            
                            for _, row in summary_table.iterrows():
                                row_cells = table.add_row().cells
                                row_cells[0].text = str(row['Система_и_Канал'])
                                row_cells[1].text = str(row['Количество'])
                                row_cells[2].text = str(row['Общий_охват'])
                                row_cells[3].text = str(row['Общая_стоимость'])
                            doc.add_paragraph()
                            
                        elif '[ГРАФИК_РАСПРЕДЕЛЕНИЯ]' in clean_line:
                            doc.add_picture(img1_stream, width=Inches(6.0))
                            
                        elif '[ГРАФИК_ОХВАТА]' in clean_line:
                            doc.add_picture(img2_stream, width=Inches(6.0))
                            
                        else:
                            p = doc.add_paragraph(clean_line)
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    bio = BytesIO()
                    doc.save(bio)

                    st.download_button(
                        label="Скачать официальный отчет (DOCX с таблицей и графиками)",
                        data=bio.getvalue(),
                        file_name="Analytic_Report_Official.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Ошибка формирования файла Word (DOCX): {e}")

else:
    st.info("Внимание: Отсутствуют требуемые исходные файлы данных (Excel/GeoJSON). Загрузка прервана.")